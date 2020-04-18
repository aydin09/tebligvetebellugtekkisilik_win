from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.shared import Length
import os
from tkinter import *
import sqlite3
import tkinter.ttk as ttk
import locale

locale.setlocale(locale.LC_ALL,"")

def bilgi_girişi(event):
    liste1=liste.get(ACTIVE)

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_edilen.delete(0,END)
    teblig_edilen_gorev.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    yaziyi_gonderen_makam.delete(0,END)
    teblig_yazi_tarih.delete(0,END)
    teblig_yazi_no.delete(0,END)
    teblig_yazi_konu.delete(0,END)
    teblig_yer.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)
                       
    vt1 = sqlite3.connect(liste1+'.sql')
    im1= vt1.cursor()
    im1.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligedilenadisoyadi TEXT, tebligedileningorevi TEXT,tebligedenadisoyadi TEXT,
                                                     tebligedengorev TEXT, yazigonderenmakam TEXT, tebligyazitarih TEXT, tebligyazino TEXT, yazikonusu TEXT,
                                                     tebligyapildigiyer TEXT, tebligyapildigitarih TEXT, tebligyapildigisaat TEXT)""")
    im1.execute("""SELECT * FROM  teblig""")
    rows = im1.fetchall()
    data_str = ""
    sf = "{}{}{}{}{}{}{}{}{}{}{}{}{}"
    for row in rows:
        data_str += sf.format(row[0],row[1],row[2],row[3],row[4],row[5],row[6],row[7],row[8],row[9],row[10],row[11],row[12])

        kaymakamlik.insert(END,row[0])
        okul_adi.insert(END,row[1])
        teblig_edilen.insert(END,row[2])
        teblig_edilen_gorev.insert(END,row[3])
        teblig_eden.insert(END,row[4])
        teblig_eden_gorev.insert(END,row[5])
        yaziyi_gonderen_makam.insert(END,row[6])
        teblig_yazi_tarih.insert(END,row[7])
        teblig_yazi_no.insert(END,row[8])
        teblig_yazi_konu.insert(END,row[9])
        teblig_yer.insert(END,row[10])
        teblig_tarih.insert(END,row[11])
        teblig_saat.insert(END,row[12])
                        
def kaydet_teblig_edilen():
    kaymakamlik1=kaymakamlik.get()
    okul_adi1=okul_adi.get()
    teblig_edilen1=teblig_edilen.get()
    teblig_edilen_gorev1=teblig_edilen_gorev.get()
    teblig_eden1=teblig_eden.get()
    teblig_eden_gorev1=teblig_eden_gorev.get()
    yaziyi_gonderen_makam1=yaziyi_gonderen_makam.get()
    teblig_yazi_tarih1=teblig_yazi_tarih.get()
    teblig_yazi_no1=teblig_yazi_no.get()
    teblig_yazi_konu1=teblig_yazi_konu.get()
    teblig_yer1=teblig_yer.get()
    teblig_tarih1=teblig_tarih.get()
    teblig_saat1=teblig_saat.get()

    if kaymakamlik1=="" or okul_adi1=="" or teblig_edilen1=="" or teblig_edilen_gorev1=="" or teblig_eden1=="" or teblig_eden_gorev1=="" or yaziyi_gonderen_makam1=="" or teblig_yazi_tarih1=="" or teblig_yazi_no1=="" or teblig_yazi_konu1=="" or teblig_yer1=="" or teblig_tarih1=="" or teblig_saat1=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Bilgileri eksiksiz giriniz!').pack()

    else:    
        kaymakamlik.delete(0,END)
        okul_adi.delete(0,END)
        teblig_edilen.delete(0,END)
        teblig_edilen_gorev.delete(0,END)
        teblig_eden.delete(0,END)
        teblig_eden_gorev.delete(0,END)
        yaziyi_gonderen_makam.delete(0,END)
        teblig_yazi_tarih.delete(0,END)
        teblig_yazi_no.delete(0,END)
        teblig_yazi_konu.delete(0,END)
        teblig_yer.delete(0,END)
        teblig_tarih.delete(0,END)
        teblig_saat.delete(0,END)

        if os.path.exists(teblig_edilen1+'.sql')== False:
            vt1 = sqlite3.connect(teblig_edilen1+'.sql')
            im1= vt1.cursor()
            im1.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligedilenadisoyadi TEXT, tebligedileningorevi TEXT,tebligedenadisoyadi TEXT,
                                                     tebligedengorev TEXT, yazigonderenmakam TEXT, tebligyazitarih TEXT, tebligyazino TEXT, yazikonusu TEXT,
                                                     tebligyapildigiyer TEXT, tebligyapildigitarih TEXT, tebligyapildigisaat TEXT)""")
            im1.execute("""INSERT INTO teblig VALUES  (?,?,?,?,?,?,?,?,?,?,?,?,?)""",(kaymakamlik1, okul_adi1, teblig_edilen1, teblig_edilen_gorev1,
                                                                                      teblig_eden1, teblig_eden_gorev1, yaziyi_gonderen_makam1,
                                                                                      teblig_yazi_tarih1, teblig_yazi_no1, teblig_yazi_konu1, teblig_yer1,
                                                                                      teblig_tarih1, teblig_saat1,))
            vt1.commit()

            liste.delete(0,END)

            for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
                if i.endswith('.sql'):
                    liste.insert(END,i[0:-4])

        else:
            vt2 = sqlite3.connect(teblig_edilen1+'.sql')
            im2= vt2.cursor()
            im2.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligedilenadisoyadi TEXT, tebligedileningorevi TEXT,tebligedenadisoyadi TEXT,
                                                     tebligedengorev TEXT, yazigonderenmakam TEXT, tebligyazitarih TEXT, tebligyazino TEXT, yazikonusu TEXT,
                                                     tebligyapildigiyer TEXT, tebligyapildigitarih TEXT, tebligyapildigisaat TEXT)""")
            im2.execute("""UPDATE teblig SET kaymakamlık=?, okul=?, tebligedilenadisoyadi=?, tebligedileningorevi=?,tebligedenadisoyadi=?,
                                             tebligedengorev=?, yazigonderenmakam=?, tebligyazitarih=?, tebligyazino=?, yazikonusu=?,tebligyapildigiyer=?,
                                             tebligyapildigitarih=?, tebligyapildigisaat=?""",(kaymakamlik1, okul_adi1, teblig_edilen1, teblig_edilen_gorev1,
                                                                                               teblig_eden1, teblig_eden_gorev1, yaziyi_gonderen_makam1,
                                                                                               teblig_yazi_tarih1, teblig_yazi_no1, teblig_yazi_konu1, teblig_yer1,
                                                                                               teblig_tarih1, teblig_saat1,))
            
            vt2.commit()
       
def cikti():
    kaymakamlik1=kaymakamlik.get()
    okul_adi1=okul_adi.get()
    teblig_edilen1=teblig_edilen.get()
    teblig_edilen_gorev1=teblig_edilen_gorev.get()
    teblig_eden1=teblig_eden.get()
    teblig_eden_gorev1=teblig_eden_gorev.get()
    yaziyi_gonderen_makam1=yaziyi_gonderen_makam.get()
    teblig_yazi_tarih1=teblig_yazi_tarih.get()
    teblig_yazi_no1=teblig_yazi_no.get()
    teblig_yazi_konu1=teblig_yazi_konu.get()
    teblig_yer1=teblig_yer.get()
    teblig_tarih1=teblig_tarih.get()
    teblig_saat1=teblig_saat.get()
    
    if kaymakamlik1=="" or okul_adi1=="" or teblig_edilen1=="" or teblig_edilen_gorev1=="" or teblig_eden1=="" or teblig_eden_gorev1=="" or yaziyi_gonderen_makam1=="" or teblig_yazi_tarih1=="" or teblig_yazi_no1=="" or teblig_yazi_konu1=="" or teblig_yer1=="" or teblig_tarih1=="" or teblig_saat1=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Bilgileri eksiksiz giriniz!').pack()

    else:    
        vt1 = sqlite3.connect(teblig_edilen1+'.sql')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligedilenadisoyadi TEXT, tebligedileningorevi TEXT,tebligedenadisoyadi TEXT,
                                                     tebligedengorev TEXT, yazigonderenmakam TEXT, tebligyazitarih TEXT, tebligyazino TEXT, yazikonusu TEXT,
                                                     tebligyapildigiyer TEXT, tebligyapildigitarih TEXT, tebligyapildigisaat TEXT)""")
        im1.execute("""SELECT * FROM  teblig""")
        rows = im1.fetchall()
        data_str = ""
        sf = "{}{}{}{}{}{}{}{}{}{}{}{}{}"
        for row in rows:
            data_str += sf.format(row[0],row[1],row[2],row[3],row[4],row[5],row[6],row[7],row[8],row[9],row[10],row[11],row[12])
        vt1.commit()        
            
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        table = document.add_table(rows=1, cols=1,style = 'Table Grid')
        
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("\nT.C.\n"+row[0]+"\n"+row[1]+"\n\n"+"TEBLİĞ VE TEBELLÜĞ BELGESİ").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        table = document.add_table(rows=2, cols=3,style = 'Table Grid')
        
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("Tebliğ Edilen Yazının Kime Gönderildiği ve Görevi").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2.5)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run(teblig_edilen1)
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[1].width = Inches(2.5)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run(teblig_edilen_gorev1)
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[2].width = Inches(2.5)

        cell = table.cell(1,0)
        table.cell(1,0).paragraphs[0].add_run("Tebliğ Edenin Adı Soyadı ve Görevi").bold = True
        table.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2.5)

        cell = table.cell(1,1)
        table.cell(1,1).paragraphs[0].add_run(teblig_eden1)
        table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[1].width = Inches(2.5)

        cell = table.cell(1,2)
        table.cell(1,2).paragraphs[0].add_run(teblig_eden_gorev1)
        table.cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[2].width = Inches(2.5)

        table = document.add_table(rows=5, cols=2,style = 'Table Grid')

        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("Yazıyı Gönderen Makam").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run(yaziyi_gonderen_makam1)
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
       

        cell = table.cell(1,0)
        table.cell(1,0).paragraphs[0].add_run("Tebliğ Yazısının Tarihi").bold = True
        table.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(1,1)
        table.cell(1,1).paragraphs[0].add_run(teblig_yazi_tarih1)
        table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
       

        cell = table.cell(2,0)
        table.cell(2,0).paragraphs[0].add_run("Tebliğ Yazısının No'su").bold = True
        table.cell(2,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(2,1)
        table.cell(2,1).paragraphs[0].add_run(teblig_yazi_no1)
        table.cell(2,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
      

        cell = table.cell(3,0)
        table.cell(3,0).paragraphs[0].add_run("Yazının/Tebliğin Konusu").bold = True
        table.cell(3,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(3,1)
        table.cell(3,1).paragraphs[0].add_run(teblig_yazi_konu1)
        table.cell(3,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
       

        cell = table.cell(4,0)
        table.cell(4,0).paragraphs[0].add_run("Tebliğin Yapıldığı Yer").bold = True
        table.cell(4,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(4,1)
        table.cell(4,1).paragraphs[0].add_run(teblig_yer1)
        table.cell(4,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
   

        table = document.add_table(rows=1, cols=3,style = 'Table Grid')

        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("Tebliğin Yapıldığı Tarih ve Saat").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2.5)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run(teblig_tarih1)
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[1].width = Inches(2.5)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run(teblig_saat1)
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[2].width = Inches(2.5)

        table = document.add_table(rows=0, cols=3,style = 'Table Grid')
        row=table.add_row().cells
        row[0].paragraphs[0].add_run("\n"+"Tebliğ Eden"+"\n\n").bold = True
        row[0].paragraphs[0].add_run(teblig_eden1+"\n")
        row[0].paragraphs[0].add_run(teblig_eden_gorev1+"\n").bold = True
        row[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Inches(2.5)

        row[1].paragraphs[0].add_run("\n"+"Tebliğ Tarihi ve Saati"+"\n\n").bold = True
        row[1].paragraphs[0].add_run(teblig_tarih1+"\n"+teblig_saat1+"\n")
        row[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        table.columns[1].width = Inches(2.5)

        row[2].paragraphs[0].add_run("\n"+"Tebellüğ Eden"+"\n\n").bold = True
        row[2].paragraphs[0].add_run(teblig_edilen1+"\n")
        row[2].paragraphs[0].add_run(teblig_edilen_gorev1+"\n").bold = True
        row[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        table.columns[2].width = Inches(2.5)

        paragraph = document.add_paragraph('\n')

        vt1 = sqlite3.connect(teblig_edilen1+'.sql')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligedilenadisoyadi TEXT, tebligedileningorevi TEXT,tebligedenadisoyadi TEXT,
                                                     tebligedengorev TEXT, yazigonderenmakam TEXT, tebligyazitarih TEXT, tebligyazino TEXT, yazikonusu TEXT,
                                                     tebligyapildigiyer TEXT, tebligyapildigitarih TEXT, tebligyapildigisaat TEXT)""")
        im1.execute("""SELECT * FROM  teblig""")
        rows = im1.fetchall()
        data_str = ""
        sf = "{}{}{}{}{}{}{}{}{}{}{}{}{}"
        for row in rows:
            data_str += sf.format(row[0],row[1],row[2],row[3],row[4],row[5],row[6],row[7],row[8],row[9],row[10],row[11],row[12])
        vt1.commit()   

        table = document.add_table(rows=1, cols=1,style = 'Table Grid')
        
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("\nT.C.\n"+row[0]+"\n"+row[1]+"\n\n"+"TEBLİĞ VE TEBELLÜĞ BELGESİ").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        table = document.add_table(rows=2, cols=3,style = 'Table Grid')
        
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("Tebliğ Edilen Yazının Kime Gönderildiği ve Görevi").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2.5)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run(teblig_edilen1)
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[1].width = Inches(2.5)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run(teblig_edilen_gorev1)
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[2].width = Inches(2.5)

        cell = table.cell(1,0)
        table.cell(1,0).paragraphs[0].add_run("Tebliğ Edenin Adı Soyadı ve Görevi").bold = True
        table.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2.5)

        cell = table.cell(1,1)
        table.cell(1,1).paragraphs[0].add_run(teblig_eden1)
        table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[1].width = Inches(2.5)

        cell = table.cell(1,2)
        table.cell(1,2).paragraphs[0].add_run(teblig_eden_gorev1)
        table.cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[2].width = Inches(2.5)

        table = document.add_table(rows=5, cols=2,style = 'Table Grid')

        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("Yazıyı Gönderen Makam").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run(yaziyi_gonderen_makam1)
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
       

        cell = table.cell(1,0)
        table.cell(1,0).paragraphs[0].add_run("Tebliğ Yazısının Tarihi").bold = True
        table.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(1,1)
        table.cell(1,1).paragraphs[0].add_run(teblig_yazi_tarih1)
        table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
       

        cell = table.cell(2,0)
        table.cell(2,0).paragraphs[0].add_run("Tebliğ Yazısının No'su").bold = True
        table.cell(2,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(2,1)
        table.cell(2,1).paragraphs[0].add_run(teblig_yazi_no1)
        table.cell(2,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
      

        cell = table.cell(3,0)
        table.cell(3,0).paragraphs[0].add_run("Yazının/Tebliğin Konusu").bold = True
        table.cell(3,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(3,1)
        table.cell(3,1).paragraphs[0].add_run(teblig_yazi_konu1)
        table.cell(3,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
       

        cell = table.cell(4,0)
        table.cell(4,0).paragraphs[0].add_run("Tebliğin Yapıldığı Yer").bold = True
        table.cell(4,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Cm(4.66)

        cell = table.cell(4,1)
        table.cell(4,1).paragraphs[0].add_run(teblig_yer1)
        table.cell(4,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
   

        table = document.add_table(rows=1, cols=3,style = 'Table Grid')

        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("Tebliğin Yapıldığı Tarih ve Saat").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2.5)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run(teblig_tarih1)
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[1].width = Inches(2.5)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run(teblig_saat1)
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[2].width = Inches(2.5)

        table = document.add_table(rows=0, cols=3,style = 'Table Grid')
        row=table.add_row().cells
        row[0].paragraphs[0].add_run("\n"+"Tebliğ Eden"+"\n\n").bold = True
        row[0].paragraphs[0].add_run(teblig_eden1+"\n")
        row[0].paragraphs[0].add_run(teblig_eden_gorev1+"\n").bold = True
        row[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Inches(2.5)

        row[1].paragraphs[0].add_run("\n"+"Tebliğ Tarihi ve Saati"+"\n\n").bold = True
        row[1].paragraphs[0].add_run(teblig_tarih1+"\n"+teblig_saat1+"\n")
        row[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        table.columns[1].width = Inches(2.5)

        row[2].paragraphs[0].add_run("\n"+"Tebellüğ Eden"+"\n\n").bold = True
        row[2].paragraphs[0].add_run(teblig_edilen1+"\n")
        row[2].paragraphs[0].add_run(teblig_edilen_gorev1+"\n").bold = True
        row[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        table.columns[2].width = Inches(2.5)        
              
        document.save('tebligtebellug.docx')

        os.startfile("tebligtebellug.docx")

def sil_teblig_edilen():
    data_sil=liste.get(ACTIVE)

    os.remove(data_sil+".sql")

    liste.delete(0,END)

    for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
        if i.endswith('.sql'):
            liste.insert(END,i[0:-4])

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_edilen.delete(0,END)
    teblig_edilen_gorev.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    yaziyi_gonderen_makam.delete(0,END)
    teblig_yazi_tarih.delete(0,END)
    teblig_yazi_no.delete(0,END)
    teblig_yazi_konu.delete(0,END)
    teblig_yer.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)

root = Tk()
root.title("Tebliğ ve Tebellüğ Belgesi Programı")
root.resizable(width=FALSE ,height=FALSE)
img=PhotoImage(file='teblig.png')
root.tk.call('wm','iconphoto',root._w,img)
mainframe = ttk.Frame(root,padding='3 3 12 12')
mainframe.grid(column=0, row=0)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight =1)

kaymakamlik = ttk.Entry(mainframe, width =60)
kaymakamlik.grid(column = 2, row = 0)

okul_adi = ttk.Entry(mainframe, width =60)
okul_adi.grid(column = 2, row = 1)

teblig_edilen = ttk.Entry(mainframe, width =60)
teblig_edilen.grid(column = 2, row = 2)

teblig_edilen_gorev = ttk.Entry(mainframe, width =60)
teblig_edilen_gorev.grid(column = 2, row = 3)

teblig_eden = ttk.Entry(mainframe, width =60)
teblig_eden.grid(column = 2, row = 4)

teblig_eden_gorev = ttk.Entry(mainframe, width =60)
teblig_eden_gorev.grid(column = 2, row = 5)

yaziyi_gonderen_makam = ttk.Entry(mainframe, width =60)
yaziyi_gonderen_makam.grid(column = 2, row = 6)

teblig_yazi_tarih = ttk.Entry(mainframe, width =60)
teblig_yazi_tarih.grid(column = 2, row = 7)

teblig_yazi_no = ttk.Entry(mainframe, width =60)
teblig_yazi_no.grid(column = 2, row = 8)

teblig_yazi_konu = ttk.Entry(mainframe, width =60)
teblig_yazi_konu.grid(column = 2, row = 9)

teblig_yer = ttk.Entry(mainframe, width =60)
teblig_yer.grid(column = 2, row = 10)

teblig_tarih = ttk.Entry(mainframe, width =60)
teblig_tarih.grid(column = 2, row = 11)

teblig_saat = ttk.Entry(mainframe, width =60)
teblig_saat.grid(column = 2, row = 12)

ttk.Label(mainframe, text ='KAYMAKAMLIK ADI').grid(column = 1, row = 0, sticky=W)
ttk.Label(mainframe, text ='OKULUN ADI').grid(column = 1, row = 1, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDİLENİN ADI SOYADI').grid(column = 1, row = 2, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDİLENİN GÖREVİ').grid(column = 1, row = 3, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDENİN ADI SOYADI').grid(column = 1, row = 4, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDENİN GÖREVİ').grid(column = 1, row = 5, sticky=W)
ttk.Label(mainframe, text ='YAZIYI GÖNDEREN MAKAM').grid(column = 1, row = 6, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ YAZISININ TARİHİ').grid(column = 1, row = 7, sticky=W)
ttk.Label(mainframe, text ="TEBLİĞ YAZISININ NO'SU").grid(column = 1, row = 8, sticky=W)
ttk.Label(mainframe, text ='YAZININ/TEBLİĞİN KONUSU').grid(column = 1, row = 9, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞİN YAPILDIĞI YER').grid(column = 1, row = 10, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞİN YAPILDIĞI TARİH').grid(column = 1, row = 11, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞİN YAPILDIĞI SAAT').grid(column = 1, row = 12, sticky=W)

ttk.Label(mainframe, text ='TEBLİĞ EDİLENLER LİSTESİ').grid(column = 5, row=0)

ttk.Label(mainframe, text ='').grid(column = 3, row=31)
ttk.Label(mainframe, text ='').grid(column = 3, row=33)

liste = Listbox(mainframe,width=30)
liste.grid(column=5, row=1,rowspan=30,  sticky=(N,S,E,W))
liste.bind("<Double-Button-1>",bilgi_girişi)

kaydirma = ttk.Scrollbar(mainframe, orient="vertical",command=liste.yview)
kaydirma.grid(column=6, row=1, rowspan=30,sticky='ns')

liste.config(yscrollcommand=kaydirma.set)
kaydirma.config(command=liste.yview)

for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
    if i.endswith('.sql'):
        liste.insert(END,i[0:-4])

ttk.Button(mainframe, text='Tebliğ Edileni Kaydet/Güncelle',command= kaydet_teblig_edilen).grid(column=5, row=32)
ttk.Button(mainframe, text='Sil', command= sil_teblig_edilen).grid(column=5, row=34)
ttk.Button(mainframe, text='Tebliğ ve Tebellüğ Ön İzleme', command = cikti).grid(column=2, row=32)

kaymakamlik.focus()

root.mainloop()    
